import csv
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from pypdf import PdfReader

from pdf_writer import write_pdf


class PdfWriterTest(unittest.TestCase):
    def test_writes_pdf_from_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "notes.txt"
            source.write_text("plain text\nsecond line", encoding="utf-8")

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "text")
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("plain text", text)

    def test_writes_pdf_from_markdown(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "notes.md"
            source.write_text("# Summary\n\n- First point\n- Second point\n", encoding="utf-8")

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "markdown")
            self.assertTrue(Path(result["output_path"]).exists())
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("Summary", text)
            self.assertIn("First point", text)

    def test_writes_pdf_from_csv(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "table.csv"
            with source.open("w", encoding="utf-8", newline="") as handle:
                writer = csv.writer(handle)
                writer.writerow(["name", "count"])
                writer.writerow(["apples", "3"])

            result = write_pdf(source)

            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("table.csv", text)
            self.assertIn("apples", text)

    def test_writes_pdf_from_excel(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "workbook.xlsx"
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = "SheetA"
            worksheet.append(["city", "count"])
            worksheet.append(["Toronto", 2])
            workbook.save(source)

            result = write_pdf(source)

            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("SheetA", text)
            self.assertIn("Toronto", text)

    def test_writes_pdf_from_word(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "letter.docx"
            document = Document()
            document.add_heading("Visit Summary", level=1)
            document.add_paragraph("Patient is recovering well.")
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Metric"
            table.rows[0].cells[1].text = "Value"
            table.rows[1].cells[0].text = "Weight"
            table.rows[1].cells[1].text = "72"
            document.save(source)

            result = write_pdf(source)

            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("Visit Summary", text)
            self.assertIn("recovering well", text)
            self.assertIn("Weight", text)

    def test_writes_pdf_from_xml(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "record.xml"
            source.write_text("<root>\n  <item id=\"1\">alpha</item>\n</root>\n", encoding="utf-8")

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "xml")
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("alpha", text)

    def test_writes_pdf_from_yaml(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "config.yml"
            source.write_text("service:\n  name: demo\n  enabled: true\n", encoding="utf-8")

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "yaml")
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("service:", text)
            self.assertIn("demo", text)

    def test_writes_pdf_from_html(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "page.html"
            source.write_text(
                (
                    "<html><body><h1>Status</h1><p>Build <strong>completed</strong> "
                    "with <em>notes</em> and <a href='https://example.com'>details</a>."
                    "</p><ol><li>alpha</li><li>beta</li></ol><pre>line 1\nline 2</pre>"
                    "</body></html>"
                ),
                encoding="utf-8",
            )

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "html")
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("Status", text)
            self.assertIn("Build completed with notes and details.", text)
            self.assertIn("alpha", text)
            self.assertIn("beta", text)
            self.assertIn("line 1", text)

    def test_writes_pdf_from_html_table(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "table.html"
            source.write_text(
                (
                    "<html><body><h2>Inventory</h2><table><tr><th>Item</th><th>Qty</th></tr>"
                    "<tr><td>Beans</td><td>4</td></tr><tr><td>Milk</td><td>2</td></tr>"
                    "</table></body></html>"
                ),
                encoding="utf-8",
            )

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "html")
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("Inventory", text)
            self.assertIn("Item", text)
            self.assertIn("Qty", text)
            self.assertIn("Beans", text)
            self.assertIn("Milk", text)

    def test_writes_pdf_from_malformed_html(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "broken.html"
            source.write_text(
                (
                    "<html><body><h1>Summary<p>Open <strong>tag<ul><li>one<li>two</ul>"
                    "<table><tr><th>Name<th>Score<tr><td>Ana<td>95</table></body></html>"
                ),
                encoding="utf-8",
            )

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "html")
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("Summary", text)
            self.assertIn("Open tag", text)
            self.assertIn("one", text)
            self.assertIn("two", text)
            self.assertIn("Ana", text)

    def test_writes_pdf_from_html_with_unsupported_links(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "links.html"
            source.write_text(
                (
                    "<html><body><a href='#'>Menu</a><a href='/workers/projects'>Projects</a>"
                    "<a href='https://example.com/docs'>Docs</a></body></html>"
                ),
                encoding="utf-8",
            )

            result = write_pdf(source)

            self.assertEqual(result["source_type"], "html")
            text = self._extract_pdf_text(result["output_path"])
            self.assertIn("Menu", text)
            self.assertIn("Projects", text)
            self.assertIn("Docs", text)

    def test_rejects_unsupported_source_type(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "notes.bin"
            source.write_text("plain text", encoding="utf-8")

            with self.assertRaises(ValueError):
                write_pdf(source)

    @staticmethod
    def _extract_pdf_text(pdf_path: str | Path) -> str:
        reader = PdfReader(str(pdf_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)


if __name__ == "__main__":
    unittest.main()