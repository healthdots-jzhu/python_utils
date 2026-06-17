"""Utilities for writing PDF files from common document formats."""

from __future__ import annotations

import argparse
import csv
import json
import re
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from xml.sax.saxutils import escape

from pypdf import PdfReader


SUPPORTED_SOURCE_SUFFIXES = {
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
    ".text": "text",
    ".xml": "xml",
    ".yml": "yaml",
    ".yaml": "yaml",
    ".htm": "html",
    ".html": "html",
    ".csv": "csv",
    ".xlsx": "excel",
    ".xlsm": "excel",
    ".xltx": "excel",
    ".xltm": "excel",
    ".docx": "word",
}


def resolve_source_path(source_path: str | Path) -> Path:
    """Resolve an input path and validate the source file type."""
    path = Path(source_path).expanduser()
    if not path.is_absolute():
        path = (Path.cwd() / path).resolve()

    if not path.exists():
        raise FileNotFoundError(f"Source file not found: {path}")

    source_type = SUPPORTED_SOURCE_SUFFIXES.get(path.suffix.lower())
    if source_type is None:
        supported = ", ".join(sorted(SUPPORTED_SOURCE_SUFFIXES))
        raise ValueError(
            f"Unsupported source file type: {path.suffix or '<none>'}. Supported: {supported}"
        )

    return path


def resolve_output_path(source_path: Path, output_path: str | Path | None) -> Path:
    """Resolve the destination PDF path."""
    if output_path is None:
        return source_path.with_suffix(".pdf")

    path = Path(output_path).expanduser()
    if not path.is_absolute():
        path = (Path.cwd() / path).resolve()

    if path.suffix.lower() != ".pdf":
        raise ValueError(f"Expected output path to end with .pdf, got: {path}")

    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _load_reportlab() -> dict[str, Any]:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError(
            "Writing PDFs requires the 'reportlab' package. Install it in the active environment."
        ) from exc

    return {
        "colors": colors,
        "letter": letter,
        "ParagraphStyle": ParagraphStyle,
        "getSampleStyleSheet": getSampleStyleSheet,
        "inch": inch,
        "PageBreak": PageBreak,
        "Paragraph": Paragraph,
        "Preformatted": Preformatted,
        "SimpleDocTemplate": SimpleDocTemplate,
        "Spacer": Spacer,
        "Table": Table,
        "TableStyle": TableStyle,
    }


def _build_styles(rl: dict[str, Any]) -> dict[str, Any]:
    stylesheet = rl["getSampleStyleSheet"]()
    styles = {
        "title": stylesheet["Title"],
        "heading1": stylesheet["Heading1"],
        "heading2": stylesheet["Heading2"],
        "heading3": stylesheet["Heading3"],
        "body": stylesheet["BodyText"],
        "table_heading": rl["ParagraphStyle"](
            "TableHeading",
            parent=stylesheet["Heading2"],
            spaceAfter=10,
            spaceBefore=12,
        ),
        "code": rl["ParagraphStyle"](
            "CodeBlock",
            parent=stylesheet["BodyText"],
            fontName="Courier",
            fontSize=9,
            leading=11,
            leftIndent=10,
            rightIndent=10,
            spaceBefore=6,
            spaceAfter=6,
        ),
    }
    return styles


def _make_table(rows: list[list[str]], rl: dict[str, Any]) -> Any:
    normalized_rows = rows or [[""]]
    max_columns = max(len(row) for row in normalized_rows)
    padded_rows = [row + [""] * (max_columns - len(row)) for row in normalized_rows]
    table = rl["Table"](padded_rows, repeatRows=1)
    table.setStyle(
        rl["TableStyle"](
            [
                ("BACKGROUND", (0, 0), (-1, 0), rl["colors"].HexColor("#d9e8fb")),
                ("TEXTCOLOR", (0, 0), (-1, 0), rl["colors"].black),
                ("GRID", (0, 0), (-1, -1), 0.5, rl["colors"].HexColor("#7a8ca8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl["colors"].white, rl["colors"].HexColor("#f5f7fb")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _load_beautiful_soup() -> dict[str, Any]:
    try:
        from bs4 import BeautifulSoup, Comment, NavigableString, Tag
    except ImportError as exc:
        raise RuntimeError(
            "Reading HTML files robustly requires the 'beautifulsoup4' package. Install it in the active environment."
        ) from exc

    return {
        "BeautifulSoup": BeautifulSoup,
        "Comment": Comment,
        "NavigableString": NavigableString,
        "Tag": Tag,
    }


def _normalize_inline_html(text: str) -> str:
    compact = re.sub(r"(?:<br\s*/?>\s*){3,}", "<br/><br/>", text)
    compact = re.sub(r">\s+<", "><", compact)
    return compact.strip()


def _safe_reportlab_href(href: str | None) -> str | None:
    if not href:
        return None

    normalized = href.strip()
    if not normalized or normalized.startswith(("#", "/", "?")):
        return None
    if normalized.startswith("//"):
        return f"https:{normalized}"

    parsed = urlparse(normalized)
    if parsed.scheme.lower() in {"http", "https", "mailto", "tel", "ftp", "file"}:
        return normalized

    return None


def _serialize_html_inline(node: Any, bs4: dict[str, Any], *, preserve_whitespace: bool = False) -> str:
    comment_type = bs4["Comment"]
    string_type = bs4["NavigableString"]
    tag_type = bs4["Tag"]

    if isinstance(node, comment_type):
        return ""

    if isinstance(node, string_type):
        text = str(node)
        if preserve_whitespace:
            return escape(text)
        collapsed = re.sub(r"\s+", " ", text)
        return escape(collapsed) if collapsed.strip() else ""

    if not isinstance(node, tag_type):
        return ""

    tag_name = node.name.lower()
    if tag_name in {"script", "style"}:
        return ""
    if tag_name == "br":
        return "\n" if preserve_whitespace else "<br/>"

    child_markup = "".join(
        _serialize_html_inline(child, bs4, preserve_whitespace=preserve_whitespace)
        for child in node.children
    )

    if tag_name in {"b", "strong"}:
        return f"<b>{child_markup}</b>"
    if tag_name in {"i", "em"}:
        return f"<i>{child_markup}</i>"
    if tag_name == "u":
        return f"<u>{child_markup}</u>"
    if tag_name == "code" and not preserve_whitespace:
        return f'<font name="Courier">{child_markup}</font>'
    if tag_name == "a":
        href = _safe_reportlab_href(node.get("href"))
        if href:
            safe_href = escape(href, {'"': '&quot;'})
            label = child_markup or escape(href)
            return f'<a href="{safe_href}">{label}</a>'

    return child_markup


def _html_table_rows(table: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.find_all("tr"):
        cells = row.find_all(["th", "td"], recursive=False) or row.find_all(["th", "td"])
        normalized_row = [
            re.sub(r"[ \t]{2,}", " ", cell.get_text("\n", strip=True)).strip()
            for cell in cells
        ]
        if normalized_row:
            rows.append(normalized_row)
    return rows


def _collect_html_blocks(node: Any, blocks: list[dict[str, Any]], bs4: dict[str, Any]) -> None:
    tag_type = bs4["Tag"]

    for child in node.children:
        if not isinstance(child, tag_type):
            continue

        tag_name = child.name.lower()
        if tag_name in {"script", "style"}:
            continue
        if tag_name in {"body", "main", "header", "footer", "aside", "section", "article", "div"}:
            _collect_html_blocks(child, blocks, bs4)
            continue
        if tag_name in {"h1", "h2", "h3", "p", "blockquote"}:
            text = _normalize_inline_html(_serialize_html_inline(child, bs4))
            if text:
                blocks.append({"kind": tag_name, "text": text})
            continue
        if tag_name == "pre":
            text = _serialize_html_inline(child, bs4, preserve_whitespace=True).strip("\n")
            if text:
                blocks.append({"kind": "pre", "text": text})
            continue
        if tag_name in {"ul", "ol"}:
            _collect_html_list_blocks(child, blocks, bs4)
            continue
        if tag_name == "table":
            rows = _html_table_rows(child)
            if rows:
                blocks.append({"kind": "table", "rows": rows})
            continue

        direct_block_children = [
            grandchild
            for grandchild in child.children
            if isinstance(grandchild, tag_type)
            and grandchild.name.lower()
            in {"h1", "h2", "h3", "p", "blockquote", "pre", "ul", "ol", "table", "section", "article", "div"}
        ]
        if direct_block_children:
            _collect_html_blocks(child, blocks, bs4)
            continue

        text = _normalize_inline_html(_serialize_html_inline(child, bs4))
        if text:
            blocks.append({"kind": "body", "text": text})


def _collect_html_list_blocks(list_node: Any, blocks: list[dict[str, Any]], bs4: dict[str, Any]) -> None:
    list_type = list_node.name.lower()
    item_index = 1
    for item in list_node.find_all("li", recursive=False):
        inline_parts: list[str] = []
        nested_lists: list[Any] = []
        for child in item.children:
            if getattr(child, "name", None) and child.name.lower() in {"ul", "ol"}:
                nested_lists.append(child)
                continue
            inline_parts.append(_serialize_html_inline(child, bs4))

        text = _normalize_inline_html("".join(inline_parts))
        prefix = f"{item_index}. " if list_type == "ol" else "• "
        if text:
            blocks.append({"kind": "body", "text": f"{prefix}{text}"})
        item_index += 1

        for nested_list in nested_lists:
            _collect_html_list_blocks(nested_list, blocks, bs4)


def _story_from_text_block(
    source_path: Path,
    rl: dict[str, Any],
    styles: dict[str, Any],
    empty_message: str,
) -> list[Any]:
    text = source_path.read_text(encoding="utf-8").strip()
    story = [
        rl["Paragraph"](escape(source_path.name), styles["title"]),
        rl["Spacer"](1, 0.2 * rl["inch"]),
    ]
    if text:
        story.append(rl["Preformatted"](text, styles["code"]))
    else:
        story.append(rl["Paragraph"](empty_message, styles["body"]))
    return story


def _story_from_html(source_path: Path, rl: dict[str, Any], styles: dict[str, Any]) -> list[Any]:
    bs4 = _load_beautiful_soup()
    soup = bs4["BeautifulSoup"](source_path.read_text(encoding="utf-8"), "html.parser")
    root = soup.body or soup
    blocks: list[dict[str, Any]] = []
    _collect_html_blocks(root, blocks, bs4)

    story: list[Any] = [
        rl["Paragraph"](escape(source_path.name), styles["title"]),
        rl["Spacer"](1, 0.2 * rl["inch"]),
    ]

    for block in blocks:
        kind = block["kind"]
        if kind == "table":
            story.append(_make_table(block["rows"], rl))
            story.append(rl["Spacer"](1, 0.12 * rl["inch"]))
            continue
        text = block["text"]
        if kind == "h1":
            style = styles["heading1"]
        elif kind == "h2":
            style = styles["heading2"]
        elif kind == "h3":
            style = styles["heading3"]
        elif kind == "pre":
            story.append(rl["Preformatted"](text, styles["code"]))
            story.append(rl["Spacer"](1, 0.12 * rl["inch"]))
            continue
        else:
            style = styles["body"]

        story.append(rl["Paragraph"](text, style))
        story.append(rl["Spacer"](1, 0.12 * rl["inch"]))

    if len(story) == 2:
        story.append(rl["Paragraph"]("(empty html document)", styles["body"]))

    return story


def _story_from_markdown(source_path: Path, rl: dict[str, Any], styles: dict[str, Any]) -> list[Any]:
    story: list[Any] = []
    lines = source_path.read_text(encoding="utf-8").splitlines()
    paragraph_lines: list[str] = []
    code_lines: list[str] = []
    in_code_block = False

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = " ".join(line.strip() for line in paragraph_lines if line.strip())
        paragraph_lines.clear()
        if text:
            story.append(rl["Paragraph"](escape(text), styles["body"]))
            story.append(rl["Spacer"](1, 0.12 * rl["inch"]))

    def flush_code() -> None:
        if not code_lines:
            return
        story.append(rl["Preformatted"]("\n".join(code_lines), styles["code"]))
        story.append(rl["Spacer"](1, 0.12 * rl["inch"]))
        code_lines.clear()

    for line in lines:
        stripped = line.rstrip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code_block:
                flush_code()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(stripped)
            continue

        if not stripped:
            flush_paragraph()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            text = escape(heading_match.group(2).strip())
            style_key = "title" if level == 1 else f"heading{min(level, 3)}"
            story.append(rl["Paragraph"](text, styles[style_key]))
            story.append(rl["Spacer"](1, 0.12 * rl["inch"]))
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet_match or ordered_match:
            flush_paragraph()
            bullet_text = bullet_match.group(1) if bullet_match else ordered_match.group(1)
            story.append(rl["Paragraph"](f"• {escape(bullet_text.strip())}", styles["body"]))
            continue

        paragraph_lines.append(stripped)

    flush_paragraph()
    flush_code()

    if not story:
        story.append(rl["Paragraph"]("(empty markdown document)", styles["body"]))

    return story


def _story_from_csv(source_path: Path, rl: dict[str, Any], styles: dict[str, Any]) -> list[Any]:
    with source_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = [[str(cell) for cell in row] for row in csv.reader(handle)]

    if not rows:
        rows = [["(empty csv file)"]]

    return [
        rl["Paragraph"](escape(source_path.name), styles["title"]),
        rl["Spacer"](1, 0.2 * rl["inch"]),
        _make_table(rows, rl),
    ]


def _story_from_excel(source_path: Path, rl: dict[str, Any], styles: dict[str, Any]) -> list[Any]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError(
            "Reading Excel files requires the 'openpyxl' package. Install it in the active environment."
        ) from exc

    workbook = load_workbook(source_path, read_only=True, data_only=True)
    try:
        story: list[Any] = [
            rl["Paragraph"](escape(source_path.name), styles["title"]),
            rl["Spacer"](1, 0.2 * rl["inch"]),
        ]

        sheet_names = workbook.sheetnames
        for index, sheet_name in enumerate(sheet_names):
            worksheet = workbook[sheet_name]
            rows = [
                ["" if value is None else str(value) for value in row]
                for row in worksheet.iter_rows(values_only=True)
            ]
            if not rows:
                rows = [["(empty sheet)"]]

            story.append(rl["Paragraph"](escape(sheet_name), styles["table_heading"]))
            story.append(_make_table(rows, rl))
            if index < len(sheet_names) - 1:
                story.append(rl["PageBreak"]())

        return story
    finally:
        workbook.close()


def _story_from_word(source_path: Path, rl: dict[str, Any], styles: dict[str, Any]) -> list[Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Reading Word files requires the 'python-docx' package. Install it in the active environment."
        ) from exc

    document = Document(str(source_path))
    story: list[Any] = [
        rl["Paragraph"](escape(source_path.name), styles["title"]),
        rl["Spacer"](1, 0.2 * rl["inch"]),
    ]

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("title"):
            style = styles["title"]
        elif style_name.startswith("heading 1"):
            style = styles["heading1"]
        elif style_name.startswith("heading 2"):
            style = styles["heading2"]
        elif style_name.startswith("heading"):
            style = styles["heading3"]
        else:
            style = styles["body"]

        prefix = "• " if "list" in style_name else ""
        story.append(rl["Paragraph"](escape(f"{prefix}{text}"), style))
        story.append(rl["Spacer"](1, 0.12 * rl["inch"]))

    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        story.append(rl["Paragraph"](f"Table {table_index}", styles["table_heading"]))
        story.append(_make_table(rows, rl))

    if len(story) == 2:
        story.append(rl["Paragraph"]("(empty Word document)", styles["body"]))

    return story


def build_story(source_path: Path, rl: dict[str, Any], styles: dict[str, Any]) -> tuple[str, list[Any]]:
    """Build a reportlab story for the supported source file."""
    source_type = SUPPORTED_SOURCE_SUFFIXES[source_path.suffix.lower()]
    if source_type == "text":
        return source_type, _story_from_text_block(source_path, rl, styles, "(empty text document)")
    if source_type == "xml":
        return source_type, _story_from_text_block(source_path, rl, styles, "(empty xml document)")
    if source_type == "yaml":
        return source_type, _story_from_text_block(source_path, rl, styles, "(empty yaml document)")
    if source_type == "html":
        return source_type, _story_from_html(source_path, rl, styles)
    if source_type == "markdown":
        return source_type, _story_from_markdown(source_path, rl, styles)
    if source_type == "csv":
        return source_type, _story_from_csv(source_path, rl, styles)
    if source_type == "excel":
        return source_type, _story_from_excel(source_path, rl, styles)
    if source_type == "word":
        return source_type, _story_from_word(source_path, rl, styles)
    raise ValueError(f"Unsupported source type: {source_path.suffix}")


def write_pdf(
    source_path: str | Path,
    output_path: str | Path | None = None,
    *,
    title: str | None = None,
) -> dict[str, Any]:
    """Write a PDF from a supported source document."""
    source = resolve_source_path(source_path)
    destination = resolve_output_path(source, output_path)
    rl = _load_reportlab()
    styles = _build_styles(rl)
    source_type, story = build_story(source, rl, styles)

    document = rl["SimpleDocTemplate"](
        str(destination),
        pagesize=rl["letter"],
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
        title=title or source.stem,
    )
    document.build(story)

    generated_reader = PdfReader(str(destination))
    return {
        "source_path": str(source),
        "source_type": source_type,
        "output_path": str(destination),
        "page_count": len(generated_reader.pages),
    }


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Write a PDF from text, markup, spreadsheet, or document source files."
    )
    parser.add_argument(
        "source_path",
        help="Path to the text, Markdown, HTML, XML, YAML, CSV, Excel, or Word file",
    )
    parser.add_argument(
        "output_path",
        nargs="?",
        help="Optional output PDF path. Defaults to the source path with a .pdf extension.",
    )
    parser.add_argument("--title", help="Optional PDF title metadata")
    return parser


def main() -> None:
    parser = _build_parser()
    args = parser.parse_args()
    payload = write_pdf(args.source_path, args.output_path, title=args.title)
    print(json.dumps(payload, indent=2, ensure_ascii=True))


if __name__ == "__main__":
    main()